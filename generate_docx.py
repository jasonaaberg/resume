#!/usr/bin/env python3
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

doc = Document()

# Set narrow margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

# Styles
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10)

# Header - Name
name = doc.add_paragraph()
name_run = name.add_run("JASON AABERG")
name_run.bold = True
name_run.font.size = Pt(22)
name_run.font.color.rgb = RGBColor(44, 62, 80)
name.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Title
title = doc.add_paragraph()
title_run = title.add_run("Senior Systems Engineer | DevOps")
title_run.font.size = Pt(12)
title_run.font.color.rgb = RGBColor(52, 73, 94)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Contact
contact = doc.add_paragraph()
contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact.add_run("956-648-1737  |  aaberg@gmail.com  |  linkedin.com/in/jason-aaberg").font.size = Pt(10)

# Add horizontal line
doc.add_paragraph("_" * 95)

# Helper function for section headers
def add_section_header(text):
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(44, 62, 80)
    p.space_after = Pt(6)

# Helper for job headers
def add_job(company, location, dates, role):
    p = doc.add_paragraph()
    p.add_run(company).bold = True
    if location:
        p.add_run(f" — {location}")
    p.add_run(f"\t\t{dates}").italic = True
    
    role_p = doc.add_paragraph()
    role_run = role_p.add_run(role)
    role_run.italic = True
    role_p.space_after = Pt(2)

# PROFESSIONAL SUMMARY
add_section_header("Professional Summary")
summary = doc.add_paragraph(
    "Senior Systems Engineer with 15+ years of experience in cloud infrastructure, automation, and DevOps. "
    "Proven track record of delivering $200K+ annual cost savings through strategic migrations and infrastructure optimization. "
    "Built AI-powered automation systems that reduced helpdesk workload by 80%. "
    "Expertise in AWS, Terraform, Ansible, CI/CD, and enterprise-scale migrations."
)
summary.space_after = Pt(12)

# KEY PROJECTS
add_section_header("Key Projects & Achievements")

projects = [
    ("AI-Powered Ticket Resolution System", 
     "Engineered an intelligent ticketing solution integrating Jira API with Confluence knowledge base. "
     "Built AI engine that analyzes 100K+ historical tickets, generates resolution paths with confidence scoring, "
     "and auto-responds to tickets above 75% confidence threshold. Reduced daily helpdesk ticket volume by 80%."),
    
    ("Multi-AI Agent Server Documentation Platform",
     "Architected custom solution to document 600+ legacy servers running 15+ years of undocumented code. "
     "Deployed agents across all servers enabling WinRM over custom ports with AWS Security Group restrictions. "
     "AI agents scan directories, identify .NET/ASP applications, and auto-generate documentation."),
    
    ("Citrix VDI to AWS AppStream Migration",
     "Led first successful migration off 20-year Citrix platform. Migrated 30+ internally built applications "
     "serving 2000+ daily users to AWS AppStream. Reduced yearly licensing and server costs from $145K to $55K (62% savings)."),
    
    ("Enterprise Identity Platform Migration",
     "Led 5 major migration projects impacting 7000+ users. Migrated email and applications from Office 365 to Google Workspace. "
     "Transitioned identity management and migrated 130+ application integrations from AD/Okta to JumpCloud."),
    
    ("Server Consolidation & Linux Migration",
     "Consolidated 120+ EOL servers down to fewer than 40 through strategic migrations. "
     "Migrated critical Windows servers to Linux. Reduced annual EC2 costs from $170K to $70K (59% savings).")
]

for title, desc in projects:
    p = doc.add_paragraph()
    p.add_run(f"• {title}: ").bold = True
    p.add_run(desc)
    p.space_after = Pt(4)

# TECHNICAL SKILLS
add_section_header("Technical Skills")

skills = [
    ("Cloud Platforms", "AWS (EC2, S3, AppStream, VPC, IAM, Security Groups, CloudWatch, Route53), GCP, Azure"),
    ("IaC & Automation", "Terraform, Ansible, PowerShell, Bash, Python, Node.js"),
    ("Containers & CI/CD", "Docker, CI/CD Pipelines, GitHub Actions"),
    ("Identity & Access", "JumpCloud, Okta, Active Directory, SSO, Google Workspace Admin"),
    ("AI/ML Operations", "Multi-agent AI architectures, OpenAI API, Jira/Confluence API integrations"),
    ("Databases", "PostgreSQL, SQL Server, Database migration and dependency mapping"),
    ("Servers", "Linux (Ubuntu, RHEL, CentOS), Windows Server, VMware, AWS EC2, AppStream, Workspaces"),
    ("Networking & Security", "TCP/IP, VPN, DNS, Load Balancers, Firewalls, WAF, Security Groups, Cisco")
]

for label, value in skills:
    p = doc.add_paragraph()
    p.add_run(f"• {label}: ").bold = True
    p.add_run(value)
    p.space_after = Pt(2)

# PROFESSIONAL EXPERIENCE
add_section_header("Professional Experience")

# Metropolis
add_job("Metropolis Technologies", "", "Feb. 2023 – Present", "Senior Systems Engineer")
exp1 = doc.add_paragraph(
    "Manage and automate infrastructure for 600+ servers across hybrid cloud environment. "
    "Lead migration initiatives, automation projects, and AI-powered operational improvements. "
    "Designed AI-powered ticket resolution system reducing helpdesk volume by 80%. "
    "Built multi-agent AI platform to document legacy servers. "
    "Led Citrix VDI to AWS AppStream migration for 2000+ users achieving $90K annual savings. "
    "Spearheaded 5 enterprise migrations impacting 7000+ users. "
    "Consolidated 120+ EOL servers reducing EC2 costs by $100K annually."
)
exp1.space_after = Pt(10)

# Valley Tech
add_job("Valley Tech Services", "Edinburg, Texas", "Nov. 2021 – Present", "Cloud Engineer Contractor (Part-time)")
exp2 = doc.add_paragraph(
    "Pineapple Products: Restructured corporate cloud infrastructure. Migrated 100+ instances to Vultr, "
    "consolidated 50+ domains to Google Cloud. Performed security audit identifying 40,000+ publicly shared sensitive documents. "
    "Saved $12,000+ in monthly recurring fees. "
    "ACI Learning: Designed 30+ cloud lab environments for CompTIA Cloud+ Bootcamp used by 8,000+ students."
)
exp2.space_after = Pt(10)

# MediaScience
add_job("MediaScience Labs", "Austin, Texas", "Jan. 2021 – Nov. 2021", "Cloud Engineer")
exp3 = doc.add_paragraph(
    "Managed hybrid environment across Austin, Chicago, and NYC offices. "
    "Designed on-demand video streaming service for participant studies of major TV networks. "
    "Implemented AWS security best practices including MFA, KMS encryption, and S3 bucket policies."
)
exp3.space_after = Pt(10)

# USGB
add_job("United States Gold Bureau", "Austin, Texas", "May 2017 – March 2020", "IT Manager")
exp4 = doc.add_paragraph(
    "Managed hybrid environment with on-prem AD, Exchange, and Citrix VDI for 150+ users. "
    "Designed network and security infrastructure for new facility. "
    "Configured company-wide SSO. Implemented NIST 800-53 IT control families."
)
exp4.space_after = Pt(10)

# Observint
add_job("Observint Technologies", "Austin, Texas", "July 2011 – May 2017", "Senior System Administrator")
exp5 = doc.add_paragraph(
    "Managed hybrid environment with 100+ users, 300+ devices, AWS instances, and 1000+ domains. "
    "Designed and deployed 10+ network test labs. Deployed scalable, highly available systems on AWS."
)
exp5.space_after = Pt(10)

# Rioplex
add_job("Rioplex Wireless", "McAllen, Texas", "June 2008 – June 2011", "Support Technician")
exp6 = doc.add_paragraph(
    "Supported 400+ end users. Configured networking equipment including Cisco, Juniper, Barracuda, and Ubiquiti devices."
)
exp6.space_after = Pt(10)

# EDUCATION
add_section_header("Education")
edu = doc.add_paragraph()
edu.add_run("University of Texas Rio Grande Valley").bold = True
edu.add_run(" — Edinburg, Texas")
edu.add_run("\t\tClass of 2007").italic = True
doc.add_paragraph("Bachelor of Arts")

# Save
doc.save('/Users/jason/resume/jason_aaberg_resume.docx')
print("Created: jason_aaberg_resume.docx")
